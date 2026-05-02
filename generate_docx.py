#!/usr/bin/env python3
"""Generate all 12 Flaney Associates blog article DOCX files."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ARTICLES_DIR = "/Users/otaigbe2013/Claude Coding/Flaney_Associates/articles"

# ── Brand colours ──────────────────────────────────────────────────────────────
C_PRIMARY  = RGBColor(0x1a, 0x3a, 0x5c)   # dark navy
C_ACCENT   = RGBColor(0x2d, 0x8c, 0xf0)   # blue
C_MUTED    = RGBColor(0x66, 0x66, 0x66)   # grey
C_TEXT     = RGBColor(0x33, 0x33, 0x33)   # near-black


# ── Helper utilities ───────────────────────────────────────────────────────────
def set_run_colour(run, colour):
    run.font.color.rgb = colour

def add_horizontal_rule(doc):
    """Insert a thin blue horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2D8CF0')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading(doc, text, level=1, colour=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = colour or C_PRIMARY
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def body(doc, text, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = colour or C_TEXT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(f'"{text}"')
    run.font.italic = True
    run.font.size   = Pt(12)
    run.font.color.rgb = C_ACCENT
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_TEXT
    p.paragraph_format.space_after = Pt(4)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_TEXT
    p.paragraph_format.space_after = Pt(4)
    return p

def meta_line(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = C_MUTED
    p.paragraph_format.space_after = Pt(2)
    return p

def footer_block(doc):
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("For more information or questions, please contact the author:")
    r.font.size = Pt(10); r.font.color.rgb = C_PRIMARY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Joshua U. Otaigbe, PhD")
    r2.font.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = C_PRIMARY
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Founder & Principal Consultant, Flaney Associates")
    r3.font.size = Pt(10); r3.font.color.rgb = C_PRIMARY
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Email: info@flaneyassociates.com  |  Web: flaneyassociates.com")
    r4.font.size = Pt(10); r4.font.color.rgb = C_ACCENT
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("Schedule a free consultation at flaneyassociates.com/contact")
    r5.font.bold = True; r5.font.size = Pt(10); r5.font.color.rgb = C_ACCENT
    doc.add_paragraph()
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run("© 2026 Flaney Associates. All rights reserved. Provided for informational purposes only.")
    r6.font.size = Pt(8); r6.font.color.rgb = C_MUTED

def article_header(doc, category, title, subtitle, date, read_time):
    p_cat = doc.add_paragraph()
    r_cat = p_cat.add_run(category.upper())
    r_cat.font.bold = True; r_cat.font.size = Pt(10); r_cat.font.color.rgb = C_ACCENT
    p_cat.paragraph_format.space_after = Pt(4)

    p_title = doc.add_heading(title, level=1)
    for run in p_title.runs:
        run.font.color.rgb = C_PRIMARY
        run.font.size = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run(subtitle)
    r_sub.font.size = Pt(13); r_sub.font.color.rgb = C_MUTED
    p_sub.paragraph_format.space_after = Pt(10)

    meta_line(doc, f"By Joshua U. Otaigbe, PhD  |  {date}  |  {read_time}")
    meta_line(doc, "Flaney Associates  |  Materials Engineering & Innovation")
    add_horizontal_rule(doc)

def new_doc():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc

def save(doc, filename):
    path = os.path.join(ARTICLES_DIR, filename)
    doc.save(path)
    print(f"  Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# ORIGINAL 6 ARTICLES
# ══════════════════════════════════════════════════════════════════════════════

# ── ARTICLE 1: AEROSPACE ──────────────────────────────────────────────────────
def docx_aerospace():
    doc = new_doc()
    article_header(doc,
        "Aerospace & Defense",
        "Next-Gen Composite Materials: How Carbon Fiber\nThermoplastics Are Reshaping Aircraft Design",
        "A plain-language guide to the materials revolution making aircraft lighter, stronger, and more fuel-efficient.",
        "March 12, 2026", "6 min read")

    heading(doc, "Introduction: Why Aircraft Materials Matter to Everyone")
    body(doc, "Every time you board a commercial flight, you are trusting your safety to the materials that make up the aircraft. For decades, aluminum has been the backbone of aviation. But the aerospace industry is in the middle of a quiet revolution. A new generation of composite materials — specifically carbon fiber-reinforced thermoplastics — is changing the way aircraft are designed, built, and maintained. These materials are not just incremental improvements; they represent a fundamental shift in what is possible in flight.")
    body(doc, "For passengers, this revolution means safer, more fuel-efficient flights. For airlines, it means lower operating costs. For the planet, it means significantly reduced carbon emissions.")

    heading(doc, "What Are Carbon Fiber Thermoplastics?")
    body(doc, "A composite is a material made from two or more substances that together create something stronger or lighter than either alone. Carbon fiber composites combine extremely thin carbon strands — about ten times thinner than a human hair — embedded in a plastic matrix. The key innovation is the type of plastic used. Traditional thermoset composites harden permanently when heated, like a cooked egg. Thermoplastic composites use plastics that can be softened by heat and reshaped, more like candle wax. This single difference opens enormous advantages in manufacturing speed, repairability, and recyclability.")

    heading(doc, "The Weight Advantage: Why Every Pound Matters")
    body(doc, "Every extra pound of structural weight means more fuel burned on every flight for the entire life of the aircraft. Carbon fiber thermoplastics offer weight savings of 20 to 40 percent compared to aluminium parts they replace. On a large commercial aircraft this can translate to thousands of pounds of weight reduction. The Boeing 787 Dreamliner uses about 50 percent composite materials by weight, contributing to fuel efficiency improvements of roughly 20 percent over comparable older aircraft.")
    callout(doc, "The shift to thermoplastic composites is not just about making lighter parts. It is about reimagining how we build aircraft entirely.")

    heading(doc, "Faster Manufacturing, Lower Costs")
    body(doc, "Thermoset composites can need hours of autoclave cure time. Thermoplastic composites can be formed in minutes using stamp forming and automated tape laying, enabling the higher production rates the industry demands.")
    heading(doc, "Key Advantages of Thermoplastic Manufacturing", level=2)
    for b in ["Production cycle times reduced from hours to minutes",
              "Parts can be welded together, eliminating thousands of fasteners",
              "Scrap material can be remelted and reused",
              "Lower energy consumption during manufacturing",
              "Easier to automate, improving consistency and quality"]:
        bullet(doc, b)

    heading(doc, "Meeting Safety Standards")
    body(doc, "The FAA and its international counterparts require extensive proof that any new material can withstand temperature swings, vibration, lightning strikes, bird impacts, and decades of pressurisation cycles. Carbon fiber thermoplastics meet these standards impressively — their damage tolerance is often superior to thermoset composites, and damage tends to remain visible rather than propagating invisibly beneath the surface.")

    heading(doc, "The Sustainability Factor")
    body(doc, "Aviation accounts for roughly 2.5 percent of global CO₂ emissions. Lighter aircraft burn less fuel. Unlike thermoset composites, thermoplastic composites can be melted down and reformed into new parts at end of service life — aligning with IATA's net-zero commitment for 2050.")

    heading(doc, "What This Means for the Future")
    body(doc, "Thermoplastic composites are currently used for secondary structures like brackets and interior panels. The technology is advancing rapidly toward primary structural applications. At Flaney Associates, we help aerospace companies navigate this transition — from material selection and process optimisation to certification support.")

    footer_block(doc)
    save(doc, "aerospace-composite-materials.docx")


# ── ARTICLE 2: AUTOMOTIVE ─────────────────────────────────────────────────────
def docx_automotive():
    doc = new_doc()
    article_header(doc,
        "Automotive",
        "The Lightweighting Imperative: How EV Manufacturers\nAre Cutting Vehicle Mass by 15%",
        "Why your next electric car will be built from a surprising mix of materials.",
        "February 28, 2026", "7 min read")

    heading(doc, "Introduction: The Range Anxiety Problem")
    body(doc, "Range anxiety remains one of the biggest barriers to EV adoption. While battery technology improves, automakers are pulling another lever to extend range: making the car lighter. Reducing vehicle weight by 10 percent can improve range by 6 to 8 percent — extra miles achieved without adding a single battery cell.")

    heading(doc, "The Multi-Material Revolution")
    body(doc, "Traditional cars were built almost entirely from steel. Today's advanced EVs use a multi-material strategy — combining several materials, each chosen for the specific job it must do.")
    heading(doc, "Materials and Their Roles", level=2)
    for row in [
        "Advanced High-Strength Steel — safety cage, structural frame — exceptional crash protection at moderate weight",
        "Aluminum Alloys — body panels, doors, hood, suspension — 40% lighter than steel with good formability",
        "Carbon Fiber Composites — roof panels, battery enclosures — extremely light and stiff",
        "Engineering Polymers — interior structures, brackets, trim — very light, design flexibility, noise reduction",
        "Magnesium Alloys — instrument panels, seat frames — lightest structural metal, 75% lighter than steel",
    ]:
        bullet(doc, row)

    heading(doc, "Real-World Results")
    body(doc, "Leading EV manufacturers have achieved 10 to 15 percent vehicle weight reductions. Aluminium-intensive body structures cut body-in-white weight by up to 40 percent. Composite battery enclosures save 25 to 30 percent weight while providing superior thermal protection.")
    callout(doc, "Every pound you remove from the vehicle is a pound you do not need to carry for 200,000 miles. The cumulative energy savings are enormous.")

    heading(doc, "The Hidden Challenge: When Materials Meet")
    body(doc, "Using multiple materials creates galvanic corrosion risks where dissimilar metals contact each other in the presence of moisture. Automakers address this through adhesive bonding, specialised coatings, and insulating fastener barriers. Poor joints can fail in just a few years.")

    heading(doc, "Crashworthiness: Lighter Does Not Mean Less Safe")
    body(doc, "Modern multi-material designs often outperform heavier predecessors in crash tests. Advanced high-strength steels are 2–3× stronger than mild steel, so thinner, lighter sheets still improve crash performance. Aluminium crumple zones absorb impact energy before it reaches the passenger compartment.")

    heading(doc, "The Cost Equation")
    body(doc, "For EVs the economics differ from conventional vehicles. Because batteries are the most expensive component, weight reduction that allows a smaller battery pack can offset the higher cost of advanced materials. Lightweighting is not just engineering — it is financial strategy.")

    heading(doc, "What This Means for the Industry")
    body(doc, "At Flaney Associates, we work with automotive manufacturers and suppliers on every aspect of the lightweighting journey — from material selection and testing to failure analysis and process optimisation.")

    footer_block(doc)
    save(doc, "automotive-lightweighting-ev.docx")


# ── ARTICLE 3: ENERGY ─────────────────────────────────────────────────────────
def docx_energy():
    doc = new_doc()
    article_header(doc,
        "Energy & Oil/Gas",
        "Corrosion-Resistant Alloys for Deepwater Pipelines:\nSelecting Materials That Survive 30+ Years Subsea",
        "A practical guide to understanding why pipeline materials matter and how the right choices prevent billion-dollar failures.",
        "February 10, 2026", "8 min read")

    heading(doc, "Introduction: The Invisible Threat Beneath the Ocean")
    body(doc, "Miles beneath the ocean surface, a vast network of pipelines carries oil and gas from the seabed to platforms. Their greatest enemy is corrosion — costing the global oil and gas industry an estimated $1.3 billion annually in failures, repairs, and lost production. The materials chosen for these pipelines are the thin line between safe operation and catastrophe.")

    heading(doc, "Understanding Subsea Corrosion")
    body(doc, "Seawater's dissolved salts, oxygen, and aggressive chemicals attack metal surfaces. Internally, the fluids often contain CO₂, H₂S, and chloride ions at high temperature and pressure — a 'sour service' environment. Pipelines must resist simultaneous attack from both sides.")
    heading(doc, "Main Types of Corrosion", level=2)
    for b in ["Uniform corrosion — gradual, even thinning of the pipe wall",
              "Pitting corrosion — localised deep attacks that can rapidly penetrate the wall",
              "Stress corrosion cracking — cracks forming under stress in a corrosive environment",
              "Sulfide stress cracking — caused by hydrogen sulfide",
              "Microbiologically influenced corrosion — accelerated by bacteria on the pipe surface"]:
        bullet(doc, b)

    heading(doc, "The Materials Toolbox")
    heading(doc, "Duplex Stainless Steels", level=2)
    body(doc, "Dual-phase microstructure gives excellent pitting and stress corrosion resistance with roughly twice the strength of standard stainless steels — the workhorse CRA for many subsea applications.")
    heading(doc, "Super Duplex Stainless Steels", level=2)
    body(doc, "Higher chromium, molybdenum, and nitrogen content for more aggressive environments. Significantly cheaper than nickel-based alternatives.")
    heading(doc, "Nickel-Based Superalloys", level=2)
    body(doc, "The most expensive option — essential for the harshest deepwater conditions where no other material will survive.")
    callout(doc, "Choosing the right alloy is not just an engineering decision. It is a risk management decision with implications measured in billions of dollars.")

    heading(doc, "Testing for a 30-Year Life")
    body(doc, "Validation requires laboratory corrosion testing in simulated service environments, full-scale pressure testing of welded pipe sections, fracture toughness testing, and hydrogen embrittlement testing. Weld qualification is equally critical — welds are often the weakest link and must demonstrate the same corrosion resistance as the parent material.")

    heading(doc, "The Economics of Material Selection")
    body(doc, "Super duplex stainless steel can cost 3–5× more per foot than carbon steel. But a carbon steel pipeline in a corrosive environment requires ongoing inhibitor injection, inspection, and eventual repair — costs that easily exceed the CRA premium over a 30-year life. The CRA premium is essentially an insurance policy against catastrophic failure.")

    heading(doc, "Looking Ahead")
    body(doc, "The transition to hydrogen as an energy carrier creates new material challenges, as hydrogen can embrittle many conventional pipeline steels. At Flaney Associates, we help energy companies make material selection decisions that balance performance, cost, and risk.")

    footer_block(doc)
    save(doc, "energy-corrosion-resistant-alloys.docx")


# ── ARTICLE 4: BIOMEDICAL ─────────────────────────────────────────────────────
def docx_biomedical():
    doc = new_doc()
    article_header(doc,
        "Biomedical",
        "Biocompatible Polymers for Implantable Devices:\nNavigating FDA Material Requirements in 2026",
        "What every medical device innovator needs to know about choosing materials that are safe for the human body and approvable by regulators.",
        "January 22, 2026", "7 min read")

    heading(doc, "Introduction: Materials That Live Inside Us")
    body(doc, "Millions of people live with implanted medical devices — hip joints, spinal cages, heart valves, pacemakers. Each must be made from materials that coexist with living tissue, sometimes for decades, without causing harm. Material selection for implantable devices requires deep understanding of both materials science and human biology, plus thorough knowledge of regulatory requirements.")

    heading(doc, "What Does Biocompatible Actually Mean?")
    body(doc, "Biocompatibility is the ability of a material to perform its intended function without causing an unacceptable adverse reaction. It is not a single property but a complex set of interactions that depends heavily on how and where the material is used. A material perfect for skin contact may be unsuitable for blood contact.")

    heading(doc, "The Polymers Leading the Way")
    heading(doc, "PEEK (Polyether Ether Ketone)", level=2)
    body(doc, "Strong, stiff, and chemically resistant with mechanical properties closer to human bone than metal. Used extensively in spinal fusion cages and orthopedic surgery, avoiding the 'stress shielding' problem of overly rigid metal implants.")
    heading(doc, "UHMWPE", level=2)
    body(doc, "Ultra-high molecular weight polyethylene has served as the bearing surface in hip and knee replacements for over 50 years. Modern cross-linked versions have dramatically reduced wear rates, extending joint replacement life.")
    heading(doc, "Bioresorbable Polymers", level=2)
    body(doc, "PLA and PGA polymers dissolve safely inside the body after serving their purpose, creating temporary scaffolds that support tissue healing and then break down into harmless natural byproducts — eliminating the need for a second surgery.")
    callout(doc, "The ideal implant material does not just avoid causing harm. It actively supports the body's natural healing processes.")

    heading(doc, "The FDA Approval Pathway")
    body(doc, "ISO 10993 defines a systematic approach to evaluating biological effects of device materials. Required tests typically include:")
    for b in ["Cytotoxicity — does the material kill or damage cells?",
              "Sensitisation — does it cause allergic reactions?",
              "Irritation — does it cause inflammation?",
              "Systemic toxicity — does it release harmful substances?",
              "Genotoxicity — does it damage DNA?",
              "Implantation testing — how does tissue respond?",
              "Hemocompatibility — how does it interact with blood?"]:
        bullet(doc, b)
    body(doc, "For permanent implants, additional chronic toxicity studies are required. The full program can take 12–24 months and cost hundreds of thousands of dollars — reason enough to make smart material choices early.")

    heading(doc, "Common Pitfalls in Material Selection")
    body(doc, "Common mistakes include selecting materials based solely on mechanical properties, underestimating sterilisation effects on polymer properties, overlooking how manufacturing can alter material behaviour, and failing to account for degradation over the implant lifetime.")

    heading(doc, "Looking Ahead")
    body(doc, "Smart polymers that respond to body conditions, 3D-printed patient-specific implants, and new surface modifications that resist bacterial colonisation are advancing rapidly. At Flaney Associates, we help medical device companies navigate from material screening through FDA submission support.")

    footer_block(doc)
    save(doc, "biomedical-biocompatible-polymers.docx")


# ── ARTICLE 5: CONSTRUCTION ───────────────────────────────────────────────────
def docx_construction():
    doc = new_doc()
    article_header(doc,
        "Construction",
        "Fiber-Reinforced Concrete: How Advanced Additives\nAre Extending Infrastructure Lifespan by Decades",
        "Why the concrete in tomorrow's bridges and buildings will be fundamentally different — and why that matters for everyone.",
        "January 8, 2026", "5 min read")

    heading(doc, "Introduction: The Infrastructure Crisis Hiding in Plain Sight")
    body(doc, "Roughly one in three U.S. bridges has been classified as structurally deficient or in need of major repair. The ASCE gives U.S. infrastructure an overall grade of C-minus. At the heart of this crisis is a materials problem: conventional concrete cracks. It is strong in compression but weak in tension. Steel rebar helps, but corrodes over time, accelerating deterioration.")

    heading(doc, "The Fiber Solution")
    body(doc, "Fiber-reinforced concrete (FRC) distributes millions of tiny fibers throughout the entire mixture. These fibers bridge across cracks as they form, preventing them from growing into the large fissures that lead to structural failure — inspired by nature, where straw fibers have reinforced adobe bricks for thousands of years.")
    heading(doc, "Types of Fibers Used in Modern Concrete", level=2)
    for b in ["Steel fibers — hooked or crimped wires providing excellent crack resistance and structural strength",
              "Glass fibers — alkali-resistant strands for architectural panels and facades",
              "Synthetic fibers — polypropylene or nylon for early-age shrinkage and fire resistance",
              "Carbon fibers — ultra-high performance for maximum strength and durability",
              "Natural fibers — cellulose-based fibers for sustainable construction"]:
        bullet(doc, b)

    heading(doc, "How Fiber Reinforcement Works")
    body(doc, "Each crack in FRC encounters thousands of bridging fibers that absorb energy and redistribute stress, slowing crack growth by up to 90 percent. Cracks stay smaller — often under 0.1 mm — enabling autogenous (self-healing) mineral sealing. The result is a material that can partially repair itself.")
    callout(doc, "Fiber-reinforced concrete does not just resist cracking. It fundamentally changes how concrete fails, turning a brittle material into one that bends before it breaks.")

    heading(doc, "Real-World Impact")
    body(doc, "Bridge decks made with FRC last 50–75 years vs 25–30 for conventional concrete. Industrial floors require 60–80 percent fewer joints. Tunnel linings show dramatically improved fire resistance. FRC costs 10–20 percent more per cubic yard but delivers 30–50 percent lower lifecycle costs when maintenance, repairs, and extended service life are factored in.")

    heading(doc, "The Self-Healing Frontier")
    body(doc, "Researchers are embedding capsules of bacteria that produce limestone, or chemicals that react with water to form mineral deposits, directly into the concrete mix. When a crack ruptures a capsule, the healing agent is released directly into the crack. Field trials on bridges and water treatment facilities are showing promising results.")

    heading(doc, "What This Means for Your Projects")
    body(doc, "FRC technology is mature, widely available, and supported by decades of field performance data. At Flaney Associates, we help construction professionals and facility owners select and specify advanced concrete systems that deliver maximum durability and long-term value.")

    footer_block(doc)
    save(doc, "construction-fiber-reinforced-concrete.docx")


# ── ARTICLE 6: CONSUMER PRODUCTS ─────────────────────────────────────────────
def docx_consumer():
    doc = new_doc()
    article_header(doc,
        "Consumer Products",
        "Sustainable Packaging Materials: Moving Beyond\nSingle-Use Plastics Without Sacrificing Performance",
        "A practical look at the materials science behind the packaging revolution and how brands can make the switch successfully.",
        "December 15, 2025", "6 min read")

    heading(doc, "Introduction: The Plastic Problem Everyone Knows About")
    body(doc, "An estimated 8 million tons of plastic waste enters the world's oceans every year, and only about 9 percent of all plastic ever produced has been recycled. Brands are responding to consumer and regulatory pressure to switch to sustainable packaging — but the materials science is complex and getting it wrong leads to product spoilage, customer complaints, and safety issues.")

    heading(doc, "Why Plastic Packaging Is So Hard to Replace")
    body(doc, "A typical food packaging film simultaneously prevents oxygen from causing spoilage, controls moisture, blocks UV light, resists punctures during shipping, seals on high-speed machines, and remains food-safe throughout shelf life. Any sustainable alternative must match most or all of these capabilities to be commercially viable.")
    callout(doc, "The challenge is not finding alternatives to plastic. The challenge is finding alternatives that work as well as plastic in the real world.")

    heading(doc, "The Sustainable Materials Landscape")
    heading(doc, "Bio-Based Polymers", level=2)
    body(doc, "Plastics made from corn starch, sugarcane, or cellulose rather than petroleum. PLA is the best-known example but has limitations: poor heat resistance, poor moisture barrier, and requires industrial composting to biodegrade. Newer PHAs offer improved properties including marine biodegradation but currently cost several times more than conventional plastics.")
    heading(doc, "Recycled-Content Materials", level=2)
    body(doc, "Post-consumer recycled PET (rPET) is already widely used in beverage bottles. Advanced chemical recycling breaks down mixed plastic waste into building blocks chemically identical to virgin plastic — keeping plastic in the economy rather than the environment.")
    heading(doc, "Paper and Fiber-Based Solutions", level=2)
    body(doc, "Paper packaging incorporating water-based or mineral barrier coatings can provide moisture and oxygen resistance needed for food packaging while remaining recyclable. Molded fiber packaging is replacing expanded polystyrene for protective packaging and food service containers.")
    heading(doc, "Compostable Films and Coatings", level=2)
    body(doc, "For food-contaminated or small flexible packaging where recycling is impractical, compostable materials break down completely in industrial or home composting systems. The key challenge is ensuring adequate composting infrastructure and clear consumer communication.")

    heading(doc, "Testing: The Critical Step Most Brands Skip")
    body(doc, "Common failures from inadequate testing include shorter shelf life (worsening food waste), package failures in shipping, incompatibility with existing packaging machinery, and consumer confusion over disposal. Rigorous shelf life, mechanical performance, equipment compatibility, and consumer disposal behaviour testing are essential.")

    heading(doc, "A Practical Framework for Making the Switch")
    for i, step in enumerate([
        "Audit your current packaging — understand exactly what each material does and why",
        "Identify the highest-impact opportunities — focus first on the most material-intensive or consumer-visible formats",
        "Evaluate alternatives rigorously — test against the full performance requirement set, not just environmental attributes",
        "Pilot before scaling — run a limited market test to identify real-world issues",
        "Communicate clearly — help consumers understand what the new packaging is made from and how to dispose of it"
    ], 1):
        numbered(doc, step)

    heading(doc, "The Business Case")
    body(doc, "60–70 percent of shoppers are willing to pay more for products in sustainable packaging. Retailers are increasingly requiring sustainability commitments from suppliers. And regulatory pressure — including EU and Canadian extended producer responsibility laws — is making brands financially responsible for packaging end-of-life management.")
    body(doc, "At Flaney Associates, we help consumer brands navigate the complex materials science behind sustainable packaging — from material evaluation and testing to supplier qualification and performance optimisation.")

    footer_block(doc)
    save(doc, "consumer-sustainable-packaging.docx")


# ══════════════════════════════════════════════════════════════════════════════
# NEW 6 ARTICLES (Round 2)
# ══════════════════════════════════════════════════════════════════════════════

# ── ARTICLE 7: AEROSPACE — Additive Manufacturing ────────────────────────────
def docx_aerospace_additive():
    doc = new_doc()
    article_header(doc,
        "Aerospace & Defense",
        "Metal 3D Printing in Aerospace: How Additive\nManufacturing Is Reinventing Aircraft Components",
        "From jet engine brackets to satellite structures, metal additive manufacturing is reshaping what is possible in aerospace design and production.",
        "April 28, 2026", "7 min read")

    heading(doc, "Introduction: Manufacturing's Quiet Revolution")
    body(doc, "For most of aviation's history, making a metal part meant machining away 80–90 percent of a solid block. Metal additive manufacturing turns this on its head — building parts layer by layer from powder or wire, adding only what is needed. This enables designs previously impossible to manufacture, dramatically reduces material waste, and opens new possibilities for rapid iteration in aerospace.")

    heading(doc, "How Metal 3D Printing Works")
    heading(doc, "Powder Bed Fusion", level=2)
    body(doc, "A thin layer of metal powder is spread across a build platform, then a high-powered laser or electron beam selectively melts the powder according to the digital design file. The process repeats layer by layer, creating precise, fully dense metal parts with internal features and lattice structures no traditional machining could produce.")
    heading(doc, "Directed Energy Deposition", level=2)
    body(doc, "A focused energy source melts metal powder or wire as it is deposited onto a substrate — useful for large structural components and for restoring worn turbine blades rather than replacing them.")

    heading(doc, "The Design Freedom Advantage")
    callout(doc, "Additive manufacturing does not just change how we make parts. It changes what parts we can imagine making.")
    heading(doc, "Topology Optimisation", level=2)
    body(doc, "Computer algorithms determine the optimal material distribution to carry specified loads — producing organic-looking designs that use material only where structurally needed. Topology-optimised titanium brackets have achieved 40–60 percent weight reductions vs conventional predecessors.")
    heading(doc, "Internal Cooling Channels", level=2)
    body(doc, "Additive manufacturing allows cooling channels with complex, optimised geometries inside hot engine components — impossible to machine. The result is better cooling with less weight and fewer parts.")
    heading(doc, "Part Consolidation", level=2)
    body(doc, "Assemblies that previously required dozens of fastened parts can often become a single printed component. GE Aviation consolidated a LEAP engine fuel nozzle from 20 parts into one — 25 percent lighter and five times more durable.")

    heading(doc, "Key Printable Aerospace Alloys")
    for row in [
        "Titanium Ti-6Al-4V — high strength-to-weight ratio — structural brackets, airframe fittings",
        "Inconel 718 — excellent high-temperature strength — turbine blades, combustors",
        "Aluminum AlSi10Mg — lightweight, good thermal properties — housings, heat exchangers",
        "17-4 PH Stainless — high strength, corrosion resistant — fittings, fasteners, tooling",
        "Cobalt-Chrome — extreme wear and heat resistance — turbine components, bearings",
    ]:
        bullet(doc, row)

    heading(doc, "Certification: The Biggest Hurdle")
    body(doc, "Aviation regulators require extensive documentation and testing before any new manufacturing process can be used for flight-critical parts. The core challenge is process variability — demonstrating that laser power, scan speed, powder characteristics, build orientation, and post-processing are sufficiently controlled to produce consistent, reliable parts.")
    heading(doc, "Key Certification Requirements", level=2)
    for b in ["Comprehensive powder feedstock characterisation and lot-to-lot variability",
              "Qualification of specific machine and parameter combinations per alloy and geometry",
              "Non-destructive inspection methods validated for additive microstructures",
              "Fatigue and fracture testing accounting for surface roughness and porosity",
              "In-process monitoring systems that detect and flag anomalies during build"]:
        bullet(doc, b)

    heading(doc, "The Road Ahead")
    body(doc, "The aerospace additive manufacturing market is growing rapidly. As certification frameworks mature, expect additive-manufactured components to move toward primary structural applications. At Flaney Associates, we help aerospace companies develop and qualify additive manufacturing processes — from alloy selection and process development to microstructural characterisation and certification support.")

    footer_block(doc)
    save(doc, "aerospace-additive-manufacturing.docx")


# ── ARTICLE 8: AUTOMOTIVE — EV Battery Materials ─────────────────────────────
def docx_automotive_battery():
    doc = new_doc()
    article_header(doc,
        "Automotive",
        "Engineering Polymers in EV Battery Systems:\nMaterials That Keep Your Battery Safe and Efficient",
        "The hidden materials science inside electric vehicle battery packs — safety, range, and longevity depend on getting it right.",
        "April 14, 2026", "7 min read")

    heading(doc, "Introduction: The Battery as a Materials Engineering Challenge")
    body(doc, "The battery pack in a modern EV is far more than a collection of cells. It must manage heat, withstand mechanical abuse, resist chemical exposure, and maintain structural integrity for hundreds of thousands of miles — all while safely containing 400–800 volts. Much of this engineering lives in the surrounding materials: polymers, composites, adhesives, and thermal interface materials that are just as critical to battery performance as the electrochemistry inside the cells.")

    heading(doc, "The Thermal Management Problem")
    body(doc, "Lithium-ion cells perform best at 15–35°C. Too cold and power drops dramatically; too hot and they degrade faster or enter thermal runaway — a chain reaction that can destroy the entire pack. Thermal interface materials (soft polymer pads placed between cells and cooling plates) must conduct heat efficiently while providing electrical insulation and accommodating cell expansion and contraction.")
    callout(doc, "In a battery pack, thermal management is not an afterthought. It is designed into every material choice from the very beginning.")

    heading(doc, "Key Polymer Applications in Battery Packs")
    heading(doc, "Cell Holders and Module Frames", level=2)
    body(doc, "Injection-moulded polymer frames must be dimensionally stable across temperature ranges, flame retardant in case of thermal runaway, and chemically resistant to electrolyte. PA66, PPS, and PBT are widely used, selected for their heat resistance and processability.")
    heading(doc, "Battery Enclosures", level=2)
    body(doc, "Carbon or glass fiber reinforced polymer enclosures can match aluminium strength at 30–40 percent less weight, contributing directly to extended range.")
    heading(doc, "Thermal Runaway Barriers", level=2)
    body(doc, "Placed between cell modules, these must contain heat, flames, and gases for at least five minutes to allow safe vehicle exit. Materials include ceramic-filled intumescent sheets, high-performance aerogel composites, and ablative coatings.")
    heading(doc, "Seals and Gaskets", level=2)
    body(doc, "Every penetration must be sealed to IP67/IP68 standards. Fluorosilicone and EPDM rubber seals remain flexible across the pack's operating range — from −40°C in arctic conditions to over 80°C during fast charging.")

    heading(doc, "The Degradation Challenge")
    body(doc, "Battery pack materials face repeated thermal cycling, vibration, mechanical shock, and potential electrolyte exposure over 8–10 years and hundreds of thousands of miles. Accelerated aging protocols — combining thermal cycling, vibration, and electrolyte immersion — are essential to validate long-term performance.")

    heading(doc, "Next-Generation Battery Chemistries and New Material Demands")
    body(doc, "Solid-state batteries replacing liquid electrolyte with solid ceramic or polymer material promise higher energy density and improved safety — but create entirely new material challenges. The solid electrolyte must maintain intimate contact with electrode materials throughout thousands of charge cycles, requiring polymer or composite components with precisely tailored mechanical compliance.")
    body(doc, "At Flaney Associates, we work with EV manufacturers and their Tier 1 and Tier 2 suppliers on the full range of battery materials challenges — from material selection and qualification testing to failure analysis of field returns.")

    footer_block(doc)
    save(doc, "automotive-ev-battery-materials.docx")


# ── ARTICLE 9: ENERGY — Renewable Materials ──────────────────────────────────
def docx_energy_renewable():
    doc = new_doc()
    article_header(doc,
        "Energy",
        "Materials for the Energy Transition: What Wind Turbines\nand Solar Panels Are Really Made Of",
        "The materials science powering the clean energy revolution — and the engineering challenges that will determine how fast it can happen.",
        "April 1, 2026", "8 min read")

    heading(doc, "Introduction: The Energy Transition Is a Materials Challenge")
    body(doc, "Solar and wind power are growing faster than any energy technology in history. But behind every panel and turbine is demanding materials engineering that rarely makes headlines. Wind turbines must survive 20–25 years in the world's harshest environments. Solar panels must maintain performance through decades of UV exposure, thermal cycling, and weather events. Getting the materials right determines whether the economics of clean energy actually work.")

    heading(doc, "Wind Turbine Blades: The Composites Challenge")
    body(doc, "A modern offshore wind turbine blade can exceed 100 metres — roughly the wingspan of an Airbus A380. It must be stiff enough to maintain aerodynamic shape, flexible enough to survive gusts, light enough to minimise structural loads, and durable enough to withstand 20 years of rain erosion, lightning, and hundreds of millions of fatigue cycles.")
    body(doc, "Glass fiber reinforced epoxy composites form most of the blade structure. Carbon fiber reinforced polymers are increasingly used for spar caps — the main load-carrying beams — because their higher stiffness enables longer blades without prohibitive weight.")
    callout(doc, "A wind turbine blade is one of the largest composite structures ever routinely manufactured. The materials science required to build them reliably is extraordinary.")

    heading(doc, "The Leading Edge Erosion Problem")
    body(doc, "Blade tips rotate at 80–100 m/s. At these velocities, raindrops impact with enormous energy, progressively eroding the aerodynamic profile — reducing energy capture by up to 5 percent annually on severely eroded blades. Protective polyurethane or epoxy coatings must balance toughness, adhesion, flexibility, and UV resistance across a 20-year service life.")

    heading(doc, "Solar Panel Materials: More Than Silicon")
    body(doc, "The silicon cell is just one layer in an engineered material stack. Each layer must perform its specific function while maintaining compatibility through thousands of thermal cycles and decades of UV radiation.")
    heading(doc, "Solar Panel Material Stack", level=2)
    for row in [
        "Cover Glass (low-iron tempered) — transmit maximum light, protect from impact and weather",
        "Front Encapsulant (EVA or POE film) — adhere to glass, protect cell from moisture, transmit light",
        "Solar Cell (monocrystalline silicon) — convert light to electricity",
        "Back Encapsulant (EVA or POE film) — protect cell rear, electrical insulation",
        "Backsheet (multi-layer TPT polymer) — final moisture barrier, electrical insulation, UV resistance",
        "Frame (anodised aluminium) — structural support, mounting, grounding",
    ]:
        bullet(doc, row)

    heading(doc, "The Encapsulant: A Critical but Overlooked Material")
    body(doc, "EVA has been the dominant encapsulant for decades but can yellow with UV exposure and release acetic acid that corrodes metal contacts. Polyolefin elastomers (POE) offer superior UV stability and moisture resistance but are harder to process and more expensive. Choosing between them requires careful consideration of climate, panel design, and project economics.")

    heading(doc, "The End-of-Life Challenge")
    body(doc, "Solar panels from the first major growth wave are beginning to reach end of service life. Developing recycling processes that efficiently separate and recover silicon, silver, and aluminium from encapsulated panels is one of the most important unsolved problems in renewable energy.")
    body(doc, "At Flaney Associates, we help energy companies across the renewable sector address materials challenges — from blade composite design and failure analysis to solar encapsulant selection and accelerated aging evaluation.")

    footer_block(doc)
    save(doc, "energy-renewable-materials.docx")


# ── ARTICLE 10: BIOMEDICAL — 3D Printed Implants ─────────────────────────────
def docx_biomedical_additive():
    doc = new_doc()
    article_header(doc,
        "Biomedical",
        "3D-Printed Implants: How Additive Manufacturing\nIs Personalizing Orthopedic Medicine",
        "Patient-specific implants printed from titanium and PEEK are transforming bone repair — here is the materials science making it possible.",
        "March 17, 2026", "6 min read")

    heading(doc, "Introduction: The Problem With Standard Sizes")
    body(doc, "Human anatomy does not come in standard sizes. Yet orthopedic implants have historically been available only in limited standard sizes and shapes, requiring surgeons to adapt procedures to fit the implant. Additive manufacturing changes this fundamentally: by combining medical imaging data with advanced materials processing, implants can now match a specific patient's anatomy with sub-millimetre precision.")

    heading(doc, "From Scan to Implant: The Workflow")
    body(doc, "The process begins with high-resolution CT scanning that creates a 3D map of the patient's anatomy. Engineers design an implant that fills the defect precisely while accounting for mechanical load requirements. The digital design is transferred to a metal additive manufacturing system that builds the implant layer by layer from titanium or cobalt-chrome powder. Turnaround from CT scan to finished implant can be a matter of days.")

    heading(doc, "Why Titanium?")
    body(doc, "Ti-6Al-4V combines excellent strength-to-weight ratio with exceptional corrosion resistance and osseointegration — the ability of bone to grow into and bond with the implant surface. But the real breakthrough is that additive manufacturing enables porous structures that mimic natural bone architecture — impossible with any conventional manufacturing method.")
    callout(doc, "A solid titanium implant and a porous printed implant of the same size can behave completely differently inside the body. The architecture is as important as the material.")

    heading(doc, "The Science of Osseointegration Through Porosity")
    body(doc, "Natural bone is hierarchically porous — from large trabecular cavities to microscopic channels for blood vessels and bone-forming cells. Smooth metal implants rely on surface chemistry and mechanical fixation. Porous additive manufactured implants provide a 3D scaffold into which bone actively grows, creating a biological bond far stronger than mechanical fixation alone. Research shows pore sizes of 300–600 µm with 65–80 percent interconnected porosity are optimal for bone ingrowth.")

    heading(doc, "Regulatory Considerations for Printed Implants")
    body(doc, "The FDA classifies patient-specific printed implants as custom devices, with a different pathway than standard manufactured implants. Critical requirements include:")
    for b in ["Validation of the design software and digital-to-physical manufacturing chain",
              "Demonstration that printed part properties meet the same standards as conventionally made implants",
              "Characterisation of residual porosity, surface roughness, and subsurface defects",
              "Verification that post-processing does not degrade critical properties",
              "Traceability from patient imaging data through design, manufacturing, and implantation"]:
        bullet(doc, b)

    heading(doc, "The Future: Bioprinting and Smart Implants")
    body(doc, "Researchers are developing bioprinting techniques where living cells are printed into 3D scaffolds to create tissues and ultimately organs. Bioprinted cartilage patches, bone scaffolds, and vascular grafts are already in early clinical trials. At Flaney Associates, we support medical device companies at the intersection of additive manufacturing, advanced materials, and regulatory compliance.")

    footer_block(doc)
    save(doc, "biomedical-3d-printed-implants.docx")


# ── ARTICLE 11: CONSTRUCTION — Smart Coatings ────────────────────────────────
def docx_construction_coatings():
    doc = new_doc()
    article_header(doc,
        "Construction",
        "Smart Coatings for Infrastructure Protection:\nHow Nanotechnology Is Defeating Corrosion",
        "The science behind protective coatings that can sense damage, respond to threats, and extend the life of steel structures by decades.",
        "March 3, 2026", "6 min read")

    heading(doc, "Introduction: The $2.5 Trillion Corrosion Problem")
    body(doc, "Corrosion costs the global economy an estimated $2.5 trillion annually — roughly 3.4 percent of global GDP. Traditional protective coatings create a physical barrier between metal and environment, but once scratched or damaged, corrosion begins and spreads beneath the intact coating, often invisibly until significant damage has occurred. A new generation of smart coatings addresses this weakness directly.")
    callout(doc, "A conventional coating waits to fail. A smart coating fights back.")

    heading(doc, "What Makes a Coating 'Smart'?")
    body(doc, "A smart coating detects a change in its environment and responds in a way that reduces damage — detecting mechanical damage or early electrochemical corrosion and releasing inhibitors in response. Think of it as a coating with a built-in immune system.")

    heading(doc, "Microencapsulated Inhibitors", level=2)
    body(doc, "Microcapsules containing corrosion inhibitors are embedded throughout the coating matrix. When a scratch ruptures the capsules, inhibitors are released directly at the damage site and react with the exposed metal surface to form a protective passivation layer. Field trials on bridges have demonstrated significantly reduced corrosion at damaged areas versus conventional coatings.")
    heading(doc, "Nanoparticle-Enhanced Barrier Coatings", level=2)
    body(doc, "Platelet-shaped nanoparticles of clay, graphene, or zinc oxide orient parallel to the coating surface, creating a tortuous path for moisture and corrosive ions. This can reduce moisture permeability by an order of magnitude compared to conventional coatings of the same thickness.")
    heading(doc, "pH-Responsive Release Systems", level=2)
    body(doc, "Corrosion changes local pH at the coating-metal interface. Hollow nanoparticles or polymer nanocontainers loaded with inhibitors release their contents only when pH drops to the level associated with active corrosion — ensuring inhibitors are available when needed but not depleted prematurely.")

    heading(doc, "Real-World Applications and Results")
    heading(doc, "Bridge and Highway Infrastructure", level=2)
    body(doc, "Early field results show that bridges coated with microcapsule-enhanced systems require recoating 30–50 percent less frequently than those using conventional coatings. For a major suspension bridge costing $5–10 million to repaint, even a 30 percent coating life extension saves millions of dollars.")
    heading(doc, "Marine and Offshore Structures", level=2)
    body(doc, "Offshore platforms, wind turbine foundations, and marine terminals face some of the world's most corrosive environments. Smart coatings that extend maintenance intervals are particularly valuable where a single painting campaign can cost millions of dollars and requires hazardous access.")

    heading(doc, "Implementation Considerations")
    body(doc, "Smart coatings are not drop-in replacements. Nanoparticles and microcapsules can be sensitive to shear forces during application, requiring modified spray settings. Field applicators need specific training to handle these materials correctly and realise their full performance potential.")
    body(doc, "At Flaney Associates, we help infrastructure owners and coating specifiers evaluate and implement smart coating technologies — from material selection and specification through application quality control and long-term performance monitoring.")

    footer_block(doc)
    save(doc, "construction-smart-coatings.docx")


# ── ARTICLE 12: CONSUMER PRODUCTS — Engineering Plastics ─────────────────────
def docx_consumer_plastics():
    doc = new_doc()
    article_header(doc,
        "Consumer Products",
        "Engineering Plastics vs. Metals: The Smart Material\nSubstitution Strategy Reshaping Product Design",
        "Why the best-designed consumer products increasingly use high-performance polymers where metal used to be the default — and how to make the switch successfully.",
        "February 17, 2026", "6 min read")

    heading(doc, "Introduction: Why Plastics Keep Winning")
    body(doc, "Pick up your smartphone. Look at a modern laptop. The best-designed consumer products have been shaped by one of manufacturing's most important trends: systematic substitution of metals with engineering-grade polymers. Not cheap commodity plastics, but high-performance materials with properties that would have seemed impossible to achieve in plastic just two decades ago.")

    heading(doc, "The Case for Engineering Plastics")
    for b in ["Weight — typically 4–7× lighter than steel and 1.5–2× lighter than aluminium",
              "Design freedom — injection moulding can produce geometries impossible to machine from metal in a single step",
              "Part consolidation — multiple metal parts can often become a single plastic component",
              "Electrical insulation — eliminates additional insulating components in electrical assemblies",
              "Corrosion resistance — no surface treatment or coating required for most applications",
              "Noise and vibration damping — superior vibration absorption vs metals",
              "Cost — typically lower part cost at high production volumes"]:
        bullet(doc, b)

    heading(doc, "A Guide to Key Engineering Plastics")
    for row in [
        "Polycarbonate (PC) — impact resistant, optically clear, heat resistant — phone cases, eyewear lenses, LED diffusers",
        "ABS — good impact strength, easy to process, paintable — appliance housings, toys, automotive trim",
        "Nylon (PA6, PA66) — strong, wear resistant, good fatigue life — gears, bearings, power tool housings",
        "POM (Acetal) — very stiff, low friction, dimensionally stable — precision gears, hinges, zippers, fasteners",
        "PEEK — extreme temperature and chemical resistance — high-end electronics, medical devices",
        "PC/ABS Blend — balance of PC toughness and ABS processability — laptop and phone housings, power tools",
    ]:
        bullet(doc, row)

    heading(doc, "When Metal Still Wins")
    body(doc, "Metals retain their advantage for very high sustained loads over long periods, elevated temperatures above most polymers' service range (generally >150–200°C), electrical and thermal conductivity, and the premium scratch-resistant feel consumers associate with high-end products.")
    callout(doc, "The best material is always the one that delivers the required performance at the lowest total cost. Sometimes that is plastic. Sometimes it is metal. Often, it is both.")

    heading(doc, "The Hidden Pitfalls of Material Substitution")
    heading(doc, "Creep and Relaxation", level=2)
    body(doc, "Unlike metals, plastics deform slowly under sustained loads even well below their short-term strength. A plastic bracket holding a sustained load may slowly deform over months until it fails. Designing for creep resistance requires understanding the long-term viscoelastic behaviour of the specific polymer under specific load conditions.")
    heading(doc, "Weld Line Weakness", level=2)
    body(doc, "When plastic flows around a core pin or through multiple mould gates, the flow fronts meet and create a weld line — a plane of weakness that can be 20–50 percent weaker than surrounding material. Critical load-bearing features must avoid weld lines or the mould must be engineered to place them in non-critical areas.")
    heading(doc, "Environmental Stress Cracking", level=2)
    body(doc, "Many engineering plastics are susceptible to stress cracking when exposed to certain chemicals, cleaning agents, or assembly lubricants while under stress — causing sudden brittle fracture at stress levels far below normal strength. Compatibility between the plastic and all chemicals encountered in manufacture, assembly, and use must be evaluated.")

    heading(doc, "A Framework for Successful Substitution")
    body(doc, "Based on our experience helping product companies make successful metal-to-plastic transitions, key steps include:")
    for step in ["Detailed load analysis — peak, sustained, and fatigue loading",
                 "Environmental analysis — temperature range, chemical exposure, UV",
                 "Material screening against the full performance envelope",
                 "Design optimisation specifically for plastic processing and load paths",
                 "Prototype testing under realistic conditions",
                 "Accelerated aging to validate long-term performance"]:
        numbered(doc, step)

    body(doc, "At Flaney Associates, we guide consumer product companies through every step of this process — from material screening and design review to failure analysis of field returns.")

    footer_block(doc)
    save(doc, "consumer-engineering-plastics.docx")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generating all 12 Flaney Associates article DOCX files...\n")
    generators = [
        # Original 6
        docx_aerospace,
        docx_automotive,
        docx_energy,
        docx_biomedical,
        docx_construction,
        docx_consumer,
        # New 6
        docx_aerospace_additive,
        docx_automotive_battery,
        docx_energy_renewable,
        docx_biomedical_additive,
        docx_construction_coatings,
        docx_consumer_plastics,
    ]
    for fn in generators:
        fn()
    print(f"\nAll 12 DOCX files saved to {ARTICLES_DIR}")
